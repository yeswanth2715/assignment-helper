from __future__ import annotations

import argparse
import json
import logging
import re
from pathlib import Path
from typing import Any

from docx import Document

try:
    import pdfplumber
except ImportError:
    pdfplumber = None  # type: ignore[assignment]

from .content_generator import build_template_content

logger = logging.getLogger(__name__)


def clean_text(text: str) -> str:
    cleaned = (
        text.replace("\xa0", " ")
        .replace("\u2019", "'")
        .replace("\u2018", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2014", "-")
        .replace("\u2013", "-")
        .replace("\u2022", "-")
        .replace("\u00ae", "(R)")
        .replace("\u2122", "(TM)")
        .replace("\u00b7", "-")
    )
    return " ".join(cleaned.split())


def markdown_escape(text: str) -> str:
    return text.replace("|", r"\|")


def extract_document(docx_path: Path) -> tuple[list[str], list[list[list[str]]]]:
    document = Document(str(docx_path))

    paragraphs = [clean_text(paragraph.text) for paragraph in document.paragraphs]
    paragraphs = [paragraph for paragraph in paragraphs if paragraph]

    tables: list[list[list[str]]] = []
    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)

    return paragraphs, tables


def extract_pdf_document(pdf_path: Path) -> tuple[list[str], list[list[list[str]]]]:
    """Extract paragraphs and tables from a PDF file using pdfplumber."""
    if pdfplumber is None:
        raise ImportError(
            "pdfplumber is required to read PDF briefs. "
            "Install it with: pip install pdfplumber>=0.10.0"
        )

    paragraphs: list[str] = []
    tables: list[list[list[str]]] = []

    with pdfplumber.open(str(pdf_path)) as pdf:
        for page in pdf.pages:
            # Extract text paragraphs
            text = page.extract_text()
            if text:
                for line in text.split("\n"):
                    cleaned = clean_text(line)
                    if cleaned:
                        paragraphs.append(cleaned)

            # Extract tables
            page_tables = page.extract_tables()
            if page_tables:
                for raw_table in page_tables:
                    rows: list[list[str]] = []
                    for raw_row in raw_table:
                        if raw_row is None:
                            continue
                        cells = [clean_text(cell) if cell else "" for cell in raw_row]
                        if any(cells):
                            rows.append(cells)
                    if rows:
                        tables.append(rows)

    return paragraphs, tables


def extract_metadata(tables: list[list[list[str]]]) -> dict[str, str]:
    metadata: dict[str, str] = {}
    if not tables:
        return metadata

    for row in tables[0]:
        if len(row) < 2:
            continue
        key = row[0].rstrip(":").strip()
        value = " ".join(cell for cell in row[1:] if cell).strip()
        if key and value:
            metadata[key] = value
    return metadata


def extract_metadata_from_paragraphs(paragraphs: list[str]) -> dict[str, str]:
    """Fallback: extract metadata from paragraphs when tables don't contain it."""
    metadata: dict[str, str] = {}
    patterns = {
        "Module": re.compile(r"(?:Module|Course)\s*:\s*(.+)", re.IGNORECASE),
        "Length": re.compile(r"(?:Word Limit|Length|Word Count)\s*:\s*(.+)", re.IGNORECASE),
        "To be submitted on": re.compile(r"(?:Submission Date|Due Date|Deadline)\s*:\s*(.+)", re.IGNORECASE),
        "Submission Method": re.compile(r"(?:Submission Method|Submit via)\s*:\s*(.+)", re.IGNORECASE),
    }
    for para in paragraphs:
        for key, pattern in patterns.items():
            if key not in metadata:
                match = pattern.search(para)
                if match:
                    metadata[key] = clean_text(match.group(1))

    # Also try structured "Key: Value" lines
    for para in paragraphs:
        kv_match = re.match(
            r"^(Assignment (?:Title|Type)|Word Limit|Weighting|Issue Date|Submission Date|Feedback Date|Issued by)\s*:\s*(.+)$",
            para,
            re.IGNORECASE,
        )
        if kv_match:
            key = clean_text(kv_match.group(1))
            value = clean_text(kv_match.group(2))
            metadata[key] = value

    # Word limit from inline pattern like "3000 words (+/- 300)"
    if "Length" not in metadata:
        for para in paragraphs:
            wl_match = re.search(r"(\d[\d,]*\s*words\s*(?:\([^)]+\))?)", para, re.IGNORECASE)
            if wl_match:
                metadata["Length"] = clean_text(wl_match.group(1))
                break

    return metadata


def collect_assignment_text(tables: list[list[list[str]]]) -> str:
    selected_chunks: list[str] = []
    for table in tables[1:5]:
        for row in table:
            chunk = " ".join(cell for cell in row if cell).strip()
            if chunk:
                selected_chunks.append(chunk)
    return clean_text(" ".join(selected_chunks))


def shorten(text: str, limit: int = 500) -> str:
    trimmed = text.strip()
    if len(trimmed) <= limit:
        return trimmed
    return trimmed[: limit - 3].rstrip() + "..."


def extract_assignment_title(paragraphs: list[str], assignment_text: str) -> str:
    title_match = re.search(
        r"Assessment Title:\s*(.+?)(?=\s+Details:|\s+TASK\b|$)",
        assignment_text,
        re.IGNORECASE,
    )
    if title_match:
        return clean_text(title_match.group(1))

    heading_match = re.search(r"#\s*([^#]+?)(?=\s*##|$)", assignment_text)
    if heading_match:
        return clean_text(heading_match.group(1))

    return paragraphs[0] if paragraphs else "Assignment Brief"


def infer_expected_output(title: str) -> str:
    lowered = title.lower()
    if "strategy" in lowered or "analysis" in lowered:
        return "Critical analysis section supported by theory and evidence."
    if "reflection" in lowered:
        return "Personal reflection section with first-person insight and applied strategy."
    if "ethical" in lowered:
        return "Issue deep-dive with an ethical framework and a stronger alternative."
    if "recommend" in lowered or "action plan" in lowered or "strateg" in lowered:
        return "Action-oriented recommendations with clear steps, owners, and measures."
    return "A structured section that answers the brief directly."


def split_section_title_and_body(block: str) -> tuple[str, str]:
    trimmed = clean_text(block)
    verb_match = re.search(
        r"\b(Map|You should|Choose|Select|Identify|Evaluate|Develop|Reflect|Construct|Apply)\b",
        trimmed,
    )
    if verb_match:
        return trimmed[: verb_match.start()].strip(), trimmed[verb_match.start() :].strip()
    return trimmed, ""


def trim_section_body(text: str) -> str:
    stop_markers = [
        "Assessment Guidelines",
        "Purpose",
        "Links to module intended learning outcomes",
        "Special Instructions",
        "Additional Assessment Components",
        "Answer the Question",
    ]
    trimmed = clean_text(text)
    stop_positions = [trimmed.find(marker) for marker in stop_markers if marker in trimmed]
    stop_positions = [position for position in stop_positions if position >= 0]
    if stop_positions:
        trimmed = trimmed[: min(stop_positions)]
    return clean_text(trimmed).lstrip(":-. ")


def extract_part_sections(assignment_text: str) -> list[dict[str, str]]:
    pattern = re.compile(
        r"(Part\s+\d+\s*-\s*.*?)(?=Part\s+\d+\s*-\s*|$)",
        re.IGNORECASE | re.DOTALL,
    )
    sections: list[dict[str, str]] = []
    for block in pattern.findall(assignment_text):
        title, body = split_section_title_and_body(block)
        clean_title = clean_text(title)
        clean_body = trim_section_body(body)
        sub_tasks = extract_sub_tasks(clean_body)
        sections.append(
            {
                "title": clean_title,
                "description": shorten(clean_body),
                "sub_tasks": sub_tasks,
                "expected_output": infer_expected_output(clean_title),
            }
        )
    return sections


def extract_named_sections(assignment_text: str, headings: list[str]) -> list[dict[str, str]]:
    lowered = assignment_text.lower()
    markers: list[tuple[int, str]] = []
    for heading in headings:
        position = lowered.find(heading.lower())
        if position >= 0:
            markers.append((position, heading))

    markers.sort()
    sections: list[dict[str, str]] = []
    for index, (start, heading) in enumerate(markers):
        end = markers[index + 1][0] if index + 1 < len(markers) else len(assignment_text)
        body = assignment_text[start + len(heading) : end]
        clean_body = trim_section_body(body)
        sub_tasks = extract_sub_tasks(clean_body)
        sections.append(
            {
                "title": heading,
                "description": shorten(clean_body),
                "sub_tasks": sub_tasks,
                "expected_output": infer_expected_output(heading),
            }
        )
    return sections


def extract_numbered_question_sections(assignment_text: str) -> list[dict[str, str]]:
    """Detect numbered questions like '1. Emotional Agility Analysis ... 2. Solutions ...'."""
    pattern = re.compile(
        r"(?:^|\s)(\d+)\.\s+([A-Z][A-Za-z ]+?)(?=\s+(?:Identify|Evaluate|Develop|Reflect|Analyze|Map|Choose|Select|Apply|Discuss|Propose|Assess|Create|Construct|Write|Explain|Compare|Recommend))",
        re.MULTILINE,
    )
    matches = list(pattern.finditer(assignment_text))
    if len(matches) < 2:
        return []

    sections: list[dict[str, str]] = []
    for index, match in enumerate(matches):
        start = match.start()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(assignment_text)
        title = clean_text(match.group(2))
        body = trim_section_body(assignment_text[match.end() : end])
        sub_tasks = extract_sub_tasks(body)
        sections.append(
            {
                "title": title,
                "description": shorten(body),
                "sub_tasks": sub_tasks,
                "expected_output": infer_expected_output(title),
            }
        )
    return sections


def detect_named_headings(assignment_text: str) -> list[str]:
    """Auto-detect capitalised headings that look like section titles."""
    heading_pattern = re.compile(
        r"(?:^|\n)\s*(?:#{1,3}\s+)?([A-Z][A-Za-z]+(?:\s+[A-Za-z]+){1,5})\s*(?:\n|$)"
    )
    candidates: list[str] = []
    for match in heading_pattern.finditer(assignment_text):
        candidate = clean_text(match.group(1))
        lowered = candidate.lower()
        skip_words = ("assignment", "assessment", "guidelines", "purpose", "instructions", "criteria", "marking", "notes")
        if not any(word in lowered for word in skip_words) and len(candidate) > 5:
            candidates.append(candidate)
    return candidates


def extract_research_paper_sections(text: str) -> list[dict[str, str]]:
    """Detect numbered research paper sections like '1. Title', '2. Abstract (150-250 words)', etc."""
    # Works on both newline-structured and flat (joined paragraph) text
    pattern = re.compile(
        r"(?:(?:^|\n|(?<=\s))|(?<=\.\s))(\d+)\.\s+((?:Title|Abstract|Introduction|Literature Review|Methodology|Results(?:\s+and\s+Data\s+Analysis)?|Data\s+Analysis|Discussion|Conclusion|Acknowledgements(?:\s+and\s+References)?|References)[^\n.]*)",
        re.IGNORECASE,
    )
    matches = list(pattern.finditer(text))
    if len(matches) < 3:
        return []

    sections: list[dict[str, str]] = []
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        raw_title = clean_text(match.group(2).split("(")[0])  # Strip word-count hints like "(150-250 words)"
        # Strip directives that bleed into the title (e.g. "Title The title should be...")
        raw_title = re.sub(r"\s+-\s+.*$", "", raw_title).strip()
        # Limit title to first 6 words maximum
        title_words = raw_title.split()
        title = " ".join(title_words[:6]).rstrip(":")
        body = trim_section_body(text[match.end() : end])

        # Extract word-count hint if present
        wc_hint = ""
        wc_match = re.search(r"\((\d+-\d+\s*words)\)", match.group(2), re.IGNORECASE)
        if wc_match:
            wc_hint = f" [{wc_match.group(1)}]"

        sub_tasks = extract_sub_tasks(body)
        sections.append(
            {
                "title": title.strip(),
                "description": shorten(body) + wc_hint,
                "sub_tasks": sub_tasks,
                "expected_output": infer_expected_output(title),
            }
        )
    return sections


def extract_task_sections(assignment_text: str) -> list[dict[str, str]]:
    part_sections = extract_part_sections(assignment_text)
    if part_sections:
        return part_sections

    # Try research paper structure (numbered: Title, Abstract, Introduction, ...)
    research_sections = extract_research_paper_sections(assignment_text)
    if research_sections:
        return research_sections

    # Try explicit named headings from the brief
    known_headings = detect_named_headings(assignment_text)
    if known_headings:
        named_sections = extract_named_sections(assignment_text, known_headings)
        if named_sections:
            return named_sections

    # Fallback: known M521-style headings
    named_sections = extract_named_sections(
        assignment_text,
        [
            "Emotional Agility Analysis",
            "Solutions and Strategies",
            "Personal Reflection",
        ],
    )
    if named_sections:
        return named_sections

    # Fallback: numbered questions
    numbered = extract_numbered_question_sections(assignment_text)
    if numbered:
        return numbered

    logger.warning("Could not detect task sections; using single catch-all section.")
    return [
        {
            "title": "Assignment Brief",
            "description": shorten(assignment_text),
            "expected_output": "A complete response aligned to the brief and rubric.",
        }
    ]


def extract_sub_tasks(description: str) -> list[str]:
    """Split a section description into discrete sub-task statements.

    Detects boundaries at:
    - Sentence breaks (period + capital letter)
    - Action-verb transitions without a period
    - 'and finally', 'then', or comma-separated action phrases.
    """
    if not description or len(description) < 30:
        return []

    action_verbs = (
        r'(?:(?:You\s+)?(?:should\s+)?'
        r'(?:Map|Choose|Select|Identify|Evaluate|Develop|Reflect|Construct|Apply|'
        r'Critically\s+evaluate|Analyze|Assess|Propose|Recommend|Compare|Discuss|'
        r'Write|Explain|Create|Consider|Examine|Investigate|Describe|Summarize|'
        r'Outline|Define|Illustrate|Contrast|Justify|Argue|Synthesize))'
    )

    # Split on: sentence boundaries, ') Verb', ', and finally Verb', ', then Verb'
    # No global re.IGNORECASE: [A-Z] in the sentence-boundary branch stays
    # case-sensitive (avoids false splits like "vs. global").  Use inline
    # (?i) so action_verb matching is case-insensitive.
    parts = re.split(
        r'(?:'
        r'(?<=\.)\s+(?=[A-Z][a-z])'
        r'|(?i:(?<=\))\s+(?=' + action_verbs + r'))'
        r'|(?i:,\s+(?:and\s+)?finally\s+(?=' + action_verbs + r'))'
        r'|(?i:,\s+then\s+(?=' + action_verbs + r'))'
        r')',
        description,
    )

    tasks: list[str] = []
    for part in parts:
        part = part.strip().strip(",;.")
        if part and len(part) > 10:
            # Re-attach leading lowercase verb swallowed by the split
            lowered = part.lower()
            leading_verbs = (
                "map ", "choose ", "select ", "identify ", "evaluate ",
                "develop ", "reflect ", "construct ", "apply ", "critically evaluate ",
                "analyze ", "assess ", "propose ", "recommend ", "compare ", "discuss ",
            )
            for v in leading_verbs:
                if lowered.startswith(v):
                    part = part[0].upper() + part[1:]
                    break
            tasks.append(part)

    # If splitting found nothing useful, return empty list
    if len(tasks) < 2:
        return []
    return tasks


def extract_weighted_criteria(tables: list[list[list[str]]]) -> list[dict[str, str]]:
    criteria: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()

    rubric_texts: list[str] = []
    for table in tables:
        table_text = clean_text(" ".join(" ".join(row) for row in table))
        lowered = table_text.lower()
        if "assessment criteria" in lowered or "following criteria" in lowered:
            rubric_texts.append(table_text)

    for rubric_text in rubric_texts:
        matches = re.findall(
            r"([A-Z][A-Za-z/'().,& -]{5,}?)\s*(?:-|:)?\s*(\d+\s*(?:%|marks))",
            rubric_text,
        )
        for criterion, weight in matches:
            clean_criterion = clean_text(criterion).rstrip(":")
            lowered = clean_criterion.lower()
            if lowered.startswith(("mark weight", "fail", "sufficient", "satisfactory", "good", "very good")):
                continue
            if "assessment criteria" in lowered:
                continue
            if "primary assessment task" in lowered:
                continue
            key = (clean_criterion, clean_text(weight))
            if key in seen:
                continue
            seen.add(key)
            criteria.append({"criterion": clean_criterion, "weight": clean_text(weight)})

    return criteria


def extract_grading_criteria_from_table(tables: list[list[list[str]]]) -> list[dict[str, str]]:
    """Extract rubric criteria from multi-column grading tables (no explicit mark weights)."""
    criteria: list[dict[str, str]] = []
    for table in tables:
        table_text = clean_text(" ".join(" ".join(row) for row in table)).lower()
        if "grading" not in table_text and "criteria" not in table_text:
            continue
        for row in table:
            if not row:
                continue
            criterion_name = clean_text(row[0])
            lowered = criterion_name.lower()
            if not criterion_name or len(criterion_name) < 5:
                continue
            if "grading" in lowered or "criteria" in lowered:
                continue
            # The remaining columns are band descriptors, not weights
            criteria.append(
                {"criterion": criterion_name, "weight": "equally weighted"}
            )
    return criteria


def infer_assignment_type(metadata: dict[str, str], assignment_text: str) -> str:
    for key in ("Primary Assessment Task", "Primary Assessment Title", "Assessment Topic"):
        if key in metadata:
            return metadata[key]

    lowered = assignment_text.lower()
    if "individual report" in lowered:
        return "Individual report"
    if "individual essay" in lowered:
        return "Individual essay"
    return "Assignment"


def build_brief_extract_markdown(
    source_file: Path,
    paragraphs: list[str],
    tables: list[list[list[str]]],
) -> str:
    lines = [
        f"# Brief Extract: {source_file.name}",
        "",
        "## Paragraphs",
        "",
    ]

    for index, paragraph in enumerate(paragraphs, start=1):
        lines.append(f"{index}. {paragraph}")

    lines.extend(["", "## Tables", ""])
    for table_index, table in enumerate(tables, start=1):
        lines.append(f"### Table {table_index}")
        lines.append("")
        column_count = max(len(row) for row in table)
        header = [f"Column {idx}" for idx in range(1, column_count + 1)]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("|" + "|".join([" --- " for _ in header]) + "|")
        for row in table:
            padded_row = row + [""] * (column_count - len(row))
            lines.append("| " + " | ".join(markdown_escape(cell) for cell in padded_row) + " |")
        lines.append("")

    return "\n".join(lines) + "\n"


def recommend_sources(section_title: str) -> str:
    lowered = section_title.lower()
    if "csr" in lowered:
        return "Company sustainability/CSR reports, Carroll (1991), Porter and Kramer (2011), stakeholder theory literature, and competitor CSR disclosures."
    if "esg" in lowered:
        return "ESG rating agencies (MSCI, Sustainalytics), GRI/SASB materiality standards, company sustainability reports, and competitor ESG benchmarks."
    if "ethical" in lowered:
        return "Applied ethics literature (Crane and Matten, 2019), NGO/watchdog reports, case-specific media coverage, and company policy responses."
    if "reflection" in lowered:
        return "Module theory (e.g. David, 2016 on emotional agility), reflective practice frameworks (Gibbs, Kolb), and personal incident notes."
    if "emotional" in lowered or "agility" in lowered:
        return "David (2016) Emotional Agility, Goleman (1995) Emotional Intelligence, Hayes et al. (2006) ACT, and leadership transition literature."
    if "solution" in lowered or "strategy" in lowered or "recommend" in lowered:
        return "Change management frameworks, applied leadership literature, feasibility evidence, and action-research methodology."
    return "Academic peer-reviewed journals, professional/industry reports, and relevant comparator case studies."


def recommend_artifact(section_title: str) -> str:
    lowered = section_title.lower()
    if "csr" in lowered:
        return "CSR initiative-to-theory mapping table or evidence matrix."
    if "esg" in lowered:
        return "Benchmark table plus one materiality or comparison visual."
    if "ethical" in lowered:
        return "Issue assessment table with gap analysis."
    if "reflection" in lowered:
        return "Short reflective section with named hooks and counter-practices."
    if "strategy" in lowered:
        return "Action plan table with timing and success measures."
    return "Argument map or evidence table."


def extract_learning_outcomes(tables: list[list[list[str]]]) -> list[str]:
    """Extract learning outcomes (LO1, LO2, ...) from brief tables."""
    outcomes: list[str] = []
    for table in tables:
        for row in table:
            text = " ".join(cell for cell in row if cell)
            lo_matches = re.findall(
                r"(LO\d+):\s*([^.]+\.)",
                text,
                re.IGNORECASE,
            )
            for lo_id, lo_text in lo_matches:
                entry = f"{lo_id.upper()}: {clean_text(lo_text)}"
                if entry not in outcomes:
                    outcomes.append(entry)
    return outcomes


def extract_los_from_paragraphs(paragraphs: list[str]) -> list[str]:
    """Fallback: extract LOs from paragraph text when they're not in tables."""
    outcomes: list[str] = []
    # Scan individual paragraphs AND the full joined text (handles multi-line LOs in PDFs)
    candidates = list(paragraphs) + [" ".join(paragraphs)]
    seen: set[str] = set()
    for text in candidates:
        lo_matches = re.findall(
            r"(LO\d+)\s*:?\s+([A-Z][^.]+\.)",
            text,
        )
        for lo_id, lo_text in lo_matches:
            entry = f"{lo_id.upper()}: {clean_text(lo_text)}"
            if entry not in seen:
                seen.add(entry)
                outcomes.append(entry)
    return outcomes


def word_count_budget(
    word_limit: str,
    rubric_criteria: list[dict[str, str]],
    num_sections: int,
) -> list[tuple[str, int]]:
    """Estimate a per-section word budget from rubric weights and total word count."""
    match = re.search(r"([\d,]+)\s*words", word_limit, re.IGNORECASE)
    if not match:
        return []
    total = int(match.group(1).replace(",", ""))

    # Parse numeric weights from rubric criteria
    weights: list[tuple[str, float]] = []
    for criterion in rubric_criteria:
        w_match = re.search(r"(\d+)", criterion.get("weight", ""))
        if w_match:
            weights.append((criterion["criterion"], float(w_match.group(1))))

    if not weights:
        # Distribute evenly across sections + intro + conclusion
        per_section = total // (num_sections + 2)
        budget = [("Introduction", per_section)]
        budget += [(f"Section {i+1}", per_section) for i in range(num_sections)]
        budget.append(("Conclusion", per_section))
        return budget

    total_weight = sum(w for _, w in weights)
    if total_weight == 0:
        return []

    budget: list[tuple[str, int]] = []
    # Reserve ~10% for intro + conclusion
    body_words = int(total * 0.9)
    budget.append(("Introduction", int(total * 0.05)))
    for name, weight in weights:
        budget.append((name, int(body_words * weight / total_weight)))
    budget.append(("Conclusion", int(total * 0.05)))
    return budget


def build_research_plan(summary: dict[str, object]) -> str:
    lines = [
        f"# Research Plan: {summary['title']}",
        "",
        "## Task-to-Evidence Matrix",
        "",
        "| Task section | What must be answered | Evidence to collect | Suggested artifact |",
        "| --- | --- | --- | --- |",
    ]

    for section in summary["task_sections"]:
        lines.append(
            "| "
            + " | ".join(
                [
                    markdown_escape(section["title"]),
                    markdown_escape(section["description"]),
                    markdown_escape(recommend_sources(section["title"])),
                    markdown_escape(recommend_artifact(section["title"])),
                ]
            )
            + " |"
        )

    if summary["rubric_criteria"]:
        lines.extend(["", "## Rubric Signals", ""])
        for item in summary["rubric_criteria"]:
            lines.append(f"- {item['criterion']} ({item['weight']})")

    return "\n".join(lines) + "\n"


def build_agent_workflow(summary: dict[str, object]) -> str:
    lines = [
        f"# Agentic Workflow: {summary['title']}",
        "",
        "## Assignment Snapshot",
        "",
        f"- Source brief: `{summary['source_file']}`",
        f"- Module: {summary['module'] or 'Not detected'}",
        f"- Assessment type: {summary['assessment_type']}",
        f"- Due date: {summary['due_date'] or 'Not detected'}",
        f"- Word limit: {summary['word_limit'] or 'Not detected'}",
        f"- Submission method: {summary['submission_method'] or 'Not detected'}",
        "",
        "## Agent Sequence",
        "",
        "### 1. Brief Intake Agent",
        "- Goal: Read the Word brief and turn it into structured requirements.",
        "- Inputs: Original `.docx` brief.",
        "- Outputs: `01_brief_extract.md`, `02_brief_summary.json`.",
        "- Done when: metadata, tasks, and rubric signals are captured cleanly.",
        "",
        "### 2. Planning Agent",
        "- Goal: Convert the brief into a work plan that maps each task to evidence and deliverables.",
        "- Inputs: brief summary and rubric.",
        "- Outputs: `03_agent_workflow.md`, `04_research_plan.md`.",
        "- Done when: every task section has a matching evidence plan and artifact type.",
        "",
        "### 3. Research Agent",
        "- Goal: Gather the sources needed to answer the brief with critical depth.",
        "- Inputs: research plan and extracted task sections.",
        "- Outputs: source notes, comparison tables, and a citation-ready evidence base.",
        "- Done when: each claim in the future draft has a credible supporting source.",
        "",
        "### 4. Structure Agent",
        "- Goal: Build the final assignment skeleton before drafting.",
        "- Inputs: brief summary, research plan, and rubric weights.",
        "- Outputs: `05_assignment_scaffold.md`.",
        "- Done when: the outline reflects both the brief and the grading logic.",
        "",
        "### 5. Drafting Agent",
        "- Goal: Write the assignment section by section in markdown first.",
        "- Inputs: scaffold, research evidence, and required visuals or tables.",
        "- Outputs: working markdown draft and any section-level tables.",
        "- Done when: all brief tasks are answered directly and critically.",
        "",
        "### 6. QA Agent",
        "- Goal: Check alignment, word count, referencing, and evidence quality.",
        "- Inputs: draft, rubric, and checklist.",
        "- Outputs: `06_execution_checklist.md` plus revision notes.",
        "- Done when: the submission is ready to export locally as Word or PDF.",
        "",
        "## Task Sections",
        "",
    ]

    for section in summary["task_sections"]:
        lines.append(f"### {section['title']}")
        lines.append(f"- Brief focus: {section['description']}")
        lines.append(f"- Expected output: {section['expected_output']}")
        lines.append("")

    if summary["rubric_criteria"]:
        lines.extend(["## Rubric Priorities", ""])
        for criterion in summary["rubric_criteria"]:
            lines.append(f"- {criterion['criterion']} ({criterion['weight']})")
        lines.append("")

    lines.extend(
        [
            "## Public Repo Boundary",
            "",
            "- Commit the workflow docs, scripts, and generated planning artifacts.",
            "- Do not commit the original Word brief, final Word submission, or exported PDF.",
        ]
    )

    return "\n".join(lines) + "\n"


def build_assignment_scaffold(summary: dict[str, object]) -> str:
    assessment_type = str(summary["assessment_type"]).lower()
    include_exec_summary = "report" in assessment_type

    lines = [
        f"# Assignment Scaffold: {summary['title']}",
        "",
        f"**Module:** {summary['module'] or '[Fill module]'}",
        "",
        f"**Assessment type:** {summary['assessment_type']}",
        "",
    ]

    # Word-count allocation table
    budget = word_count_budget(
        str(summary.get("word_limit", "")),
        summary.get("rubric_criteria", []),
        len(summary.get("task_sections", [])),
    )
    if budget:
        lines.extend(
            [
                "## Word-Count Allocation (Indicative)",
                "",
                "| Section | Approx. words |",
                "| --- | --- |",
            ]
        )
        for name, words in budget:
            lines.append(f"| {name} | ~{words} |")
        lines.append("")

    # Learning outcomes mapping
    los = summary.get("learning_outcomes", [])
    if los:
        lines.extend(["## Learning Outcomes Addressed", ""])
        for lo in los:
            lines.append(f"- {lo}")
        lines.append("")

    if include_exec_summary:
        lines.extend(
            [
                "## Executive Summary",
                "",
                "- Summarise the central argument, strongest evidence, biggest risk, and key recommendation.",
                "",
                "## List of Tables and Figures",
                "",
                "- Add any comparison tables, issue maps, or visuals used in the final report.",
                "",
            ]
        )

    lines.extend(["## Introduction", "", "- Frame the assignment, scope, and central argument.", ""])

    for section in summary["task_sections"]:
        lines.extend([f"## {section['title']}", ""])

        sub_tasks: list[str] = section.get("sub_tasks", [])  # type: ignore[assignment]
        if sub_tasks:
            for i, task in enumerate(sub_tasks, 1):
                heading = task.strip()
                if heading and heading[0].islower():
                    heading = heading[0].upper() + heading[1:]
                lines.extend(
                    [
                        f"### {i}. {heading}",
                        "",
                        f"- Required focus: {task}",
                        "- Theory to integrate:",
                        "- Evidence to cite:",
                        "",
                    ]
                )
        else:
            lines.extend(
                [
                    f"- Required focus: {section['description']}",
                    "- Theory to integrate:",
                    "- Evidence to cite:",
                ]
            )

        lines.extend(
            [
                f"- Suggested artifact: {recommend_artifact(section['title'])}",
                "",
            ]
        )

    if any("reflection" in section["title"].lower() for section in summary["task_sections"]):
        lines.extend(["## Reflection Close", "", "- End with what changes in your own practice.", ""])

    lines.extend(
        [
            "## Conclusion",
            "",
            "- Synthesize the full answer and restate the judgement.",
            "",
            "## Reference List",
            "",
            "- Add Harvard-style references.",
            "",
        ]
    )

    return "\n".join(lines)


def build_execution_checklist(summary: dict[str, object]) -> str:
    word_limit = str(summary.get("word_limit", ""))
    match = re.search(r"([\d,]+)\s*words", word_limit, re.IGNORECASE)
    word_count_str = match.group(1) if match else "[check brief]"

    lines = [
        f"# Execution Checklist: {summary['title']}",
        "",
        "## Intake",
        "",
        "- [ ] Confirm the module, due date, submission format, and word count.",
        "- [ ] Extract all assignment parts from the Word brief.",
        "- [ ] Capture rubric criteria and weighting where available.",
        "",
        "## Planning",
        "",
        "- [ ] Turn every brief section into a draft section or subsection.",
        "- [ ] Identify the theories, frameworks, or models required.",
        "- [ ] Build a source list for both academic and practical evidence.",
        "- [ ] Map each learning outcome to the section(s) that address it.",
        "",
        "## Drafting",
        "",
    ]

    for section in summary["task_sections"]:
        lines.append(f"- [ ] Complete `{section['title']}`.")

    lines.extend(
        [
            "",
            "## Quality Gate — MSc Standard",
            "",
            f"- [ ] Word count is within {word_count_str} +/- 10% (excluding title page, contents, executive summary, references, appendices).",
            "- [ ] Every rubric criterion is demonstrably addressed in the text.",
            "- [ ] Every learning outcome (LO) is covered by at least one section.",
            "- [ ] The answer is analytical and evaluative, not just descriptive.",
            "- [ ] Recommendations or reflections follow logically from the analysis.",
            "- [ ] Each section integrates relevant academic theory (not just practice).",
            "- [ ] Harvard referencing is consistent: every in-text citation has a reference entry and vice versa.",
            "- [ ] Sources include a mix of peer-reviewed journals, professional reports, and primary data.",
            "- [ ] Visuals and tables do analytical work; each has a descriptive caption and source line.",
            "- [ ] Submission format matches the brief (Word or PDF, file naming convention).",
            "- [ ] Final draft has been proofread for grammar, spelling, and academic tone.",
            "- [ ] Turnitin/originality check completed; no unattributed content.",
            "- [ ] Export the final Word or PDF locally without committing it to Git.",
            "",
        ]
    )

    # Learning outcomes checklist
    los = summary.get("learning_outcomes", [])
    if los:
        lines.extend(["## Learning Outcome Coverage", ""])
        for lo in los:
            lines.append(f"- [ ] {lo}")
        lines.append("")

    return "\n".join(lines)


def build_draft_starter(summary: dict[str, object]) -> str:
    assessment_type = str(summary["assessment_type"]).lower()
    include_exec_summary = "report" in assessment_type

    lines = [
        f"# Draft Starter: {summary['title']}",
        "",
        "> This file is an editable starter, not a submission-ready assignment.",
        "> Replace every placeholder with your own verified analysis, evidence, and citations.",
        "",
        f"**Module:** {summary['module'] or '[Fill module]'}",
        "",
        f"**Assessment type:** {summary['assessment_type']}",
        "",
        "**Student name:** [Add your name]",
        "",
        "**Student ID:** [Add your student ID]",
        "",
        "## Working Thesis",
        "",
        "- [Write the central judgement you want the assignment to defend.]",
        "",
    ]

    if include_exec_summary:
        lines.extend(
            [
                "## Executive Summary",
                "",
                "- [Summarise the argument, strongest evidence, and main recommendation.]",
                "",
                "## List of Tables and Figures",
                "",
                "- [Add planned tables and visuals after you confirm they help the argument.]",
                "",
            ]
        )

    lines.extend(["## Introduction", "", "- [Frame the brief, scope, and argument. Add at least one verified citation.]", ""])

    for section in summary["task_sections"]:
        lines.extend([f"## {section['title']}", ""])

        sub_tasks: list[str] = section.get("sub_tasks", [])  # type: ignore[assignment]

        if sub_tasks:
            lines.append(f"**Brief requirement:** {section['description']}")
            lines.append("")
            for i, task in enumerate(sub_tasks, 1):
                heading = task.strip()
                if heading and heading[0].islower():
                    heading = heading[0].upper() + heading[1:]
                lines.extend(
                    [
                        f"### {i}. {heading}",
                        "",
                        "- [Answer this sub-task directly. Apply theory, cite evidence, and state your judgement.]",
                        "",
                    ]
                )
        else:
            lines.extend(
                [
                    f"**Brief requirement:** {section['description']}",
                    "",
                ]
            )

        lines.extend(
            [
                "**Section judgement:**",
                "",
                "- [State the main claim you will defend in this section.]",
                "",
                "**Theory or framework to apply:**",
                "",
                "- [Add the named model, theory, or standard you will use.]",
                "",
                "**Evidence to insert:**",
                "",
                "- [Source 1: what it proves]",
                "- [Source 2: what it proves]",
                "- [Source 3: what it proves]",
                "",
                "**Draft paragraph plan:**",
                "",
                "- Paragraph 1: [Set context and scope.]",
                "- Paragraph 2: [Apply theory or framework.]",
                "- Paragraph 3: [Critically evaluate the evidence.]",
                "- Paragraph 4: [State implications, comparison, or recommendation.]",
                "",
                "**In-text citation slots:**",
                "",
                "- [Example placeholder: (Author, Year)]",
                "- [Verify page numbers, years, and source names before submission.]",
                "",
                f"**Planned asset:** {recommend_artifact(section['title'])}",
                "",
            ]
        )

    if any("reflection" in section["title"].lower() for section in summary["task_sections"]):
        lines.extend(
            [
                "## Reflection Close",
                "",
                "- [State what changes in your own practice and why.]",
                "",
            ]
        )

    lines.extend(
        [
            "## Conclusion",
            "",
            "- [Synthesize the full answer and restate the judgement.]",
            "",
            "## Reference List",
            "",
            "- [Add only sources you actually used and verified.]",
            "- [Format them in Harvard style using the guide in `09_harvard_reference_guide.md`.]",
            "",
            "## Rubric Self-Audit",
            "",
        ]
    )

    rubric_criteria = summary.get("rubric_criteria", [])
    if rubric_criteria:
        for criterion in rubric_criteria:
            lines.append(f"- [ ] **{criterion['criterion']}** ({criterion['weight']}): [Note where and how you demonstrate this.]")
    else:
        lines.append("- [ ] [Check rubric criteria and confirm each is addressed in the text.]")

    lines.extend(
        [
            "",
            "## Final Check Before Submission",
            "",
            "- [ ] Every citation matches a real source.",
            "- [ ] Every claim is supported by evidence.",
            "- [ ] All placeholders have been replaced.",
            "- [ ] Word count is within the permitted range.",
            "- [ ] Submission format matches the brief requirements.",
            "- [ ] The assignment reflects your own judgement and writing.",
            "",
        ]
    )

    return "\n".join(lines)


def suggest_assets(summary: dict[str, object]) -> list[dict[str, str]]:
    suggestions: list[dict[str, str]] = [
        {
            "asset": "Table 1. Requirement-to-Argument Map",
            "use": "Near the introduction or methodology section",
            "purpose": "Show how each brief task is being answered.",
            "data": "Task sections, section claims, and planned evidence.",
        }
    ]

    for section in summary["task_sections"]:
        title = section["title"].lower()
        if "csr" in title:
            suggestions.append(
                {
                    "asset": "Table. CSR initiative-to-theory map",
                    "use": section["title"],
                    "purpose": "Link company initiatives to the selected framework clearly.",
                    "data": "CSR initiatives, chosen theory, alignment points, and critique points.",
                }
            )
        elif "esg" in title:
            suggestions.append(
                {
                    "asset": "Figure. ESG benchmark or materiality visual",
                    "use": section["title"],
                    "purpose": "Summarize comparison or materiality logic at a glance.",
                    "data": "Competitor metrics, ratings, material issues, or trend data.",
                }
            )
            suggestions.append(
                {
                    "asset": "Table. ESG comparison matrix",
                    "use": section["title"],
                    "purpose": "Compare environment, social, and governance performance clearly.",
                    "data": "Company KPIs, competitor KPIs, and critical commentary.",
                }
            )
        elif "ethical" in title:
            suggestions.append(
                {
                    "asset": "Table. Ethical issue gap analysis",
                    "use": section["title"],
                    "purpose": "Compare current response, evidence of progress, and remaining gaps.",
                    "data": "Current company actions, stakeholder concerns, and ethical framework criteria.",
                }
            )
        elif "strategy" in title or "solutions" in title or "recommend" in title:
            suggestions.append(
                {
                    "asset": "Table. Action plan",
                    "use": section["title"],
                    "purpose": "Translate recommendations into timing, owner, and success measures.",
                    "data": "Actions, rationale, timing, and metrics.",
                }
            )
        elif "reflection" in title:
            suggestions.append(
                {
                    "asset": "Optional table. Personal hook-to-practice map",
                    "use": section["title"],
                    "purpose": "Keep the reflection structured without overexplaining it.",
                    "data": "Named hook, impact, and counter-practice.",
                }
            )
        else:
            suggestions.append(
                {
                    "asset": "Table. Evidence map",
                    "use": section["title"],
                    "purpose": "Keep theory, evidence, and judgement aligned.",
                    "data": "Claims, sources, frameworks, and critique points.",
                }
            )

    return suggestions


def build_figure_table_plan(summary: dict[str, object]) -> str:
    lines = [
        f"# Figure and Table Plan: {summary['title']}",
        "",
        "Use visuals only when they clarify the argument. Every figure or table should do analytical work, not just decorate the paper.",
        "",
        "| Planned asset | Where to use it | Purpose | What data you still need |",
        "| --- | --- | --- | --- |",
    ]

    for suggestion in suggest_assets(summary):
        lines.append(
            "| "
            + " | ".join(
                [
                    markdown_escape(suggestion["asset"]),
                    markdown_escape(suggestion["use"]),
                    markdown_escape(suggestion["purpose"]),
                    markdown_escape(suggestion["data"]),
                ]
            )
            + " |"
        )

    lines.extend(
        [
            "",
            "## Caption Checklist",
            "",
            "- Give each figure and table a descriptive title.",
            "- Add a source line under every non-original asset.",
            "- If you created the visual yourself, still note the underlying data source.",
            "- Make sure the text explicitly refers to each asset and explains why it matters.",
            "",
        ]
    )

    return "\n".join(lines)


def build_harvard_reference_guide() -> str:
    lines = [
        "# Harvard Reference and In-Text Citation Guide",
        "",
        "Use this file as a formatting aide only. Verify every source against the original publication before submitting.",
        "",
        "## In-Text Citation Patterns",
        "",
        "- Narrative: `Author (Year) argues that ...`",
        "- Parenthetical: `... (Author, Year).`",
        "- Two authors: `(Author and Author, Year)`",
        "- Three or more authors: `(Author et al., Year)`",
        "- Direct quote: `(Author, Year, p. xx)`",
        "- Secondary citation: `(Author, Year, cited in Author, Year)`",
        "",
        "## Reference Templates",
        "",
        "### Journal Article",
        "",
        "`Author, A.A. and Author, B.B. (Year) 'Article title', Journal Title, volume(issue), pp. xx-xx. Available at: URL (Accessed: Date).`",
        "",
        "### Book",
        "",
        "`Author, A.A. (Year) Book Title. Place of publication: Publisher.`",
        "",
        "### Book Chapter (Edited Collection)",
        "",
        "`Author, A.A. (Year) 'Chapter title', in Editor, E.E. (ed.) Book Title. Place of publication: Publisher, pp. xx-xx.`",
        "",
        "### Conference Paper",
        "",
        "`Author, A.A. (Year) 'Paper title', Proceedings of Conference Name. Location, Date. Place of publication: Publisher, pp. xx-xx.`",
        "",
        "### Company Report",
        "",
        "`Company Name (Year) Report title. Available at: URL (Accessed: Date).`",
        "",
        "### Government or Policy Report",
        "",
        "`Department/Agency Name (Year) Report title. Place of publication: Publisher. Available at: URL (Accessed: Date).`",
        "",
        "### Newspaper Article",
        "",
        "`Author, A.A. (Year) 'Article title', Newspaper Title, Date, p. xx. Available at: URL (Accessed: Date).`",
        "",
        "### Web Page or NGO Report",
        "",
        "`Organisation Name (Year) Page or report title. Available at: URL (Accessed: Date).`",
        "",
        "## Quality Checks",
        "",
        "- Make sure every in-text citation has a matching reference entry.",
        "- Make sure every reference entry is cited in the text.",
        "- Do not invent page numbers, authors, or publication years.",
        "- Use the same naming convention for the same source throughout.",
        "- Aim for a balance of peer-reviewed and professional sources.",
        "- Secondary citations should be used sparingly; always try to access the original source.",
        "",
    ]

    return "\n".join(lines)


def build_summary(
    source_file: Path,
    paragraphs: list[str],
    tables: list[list[list[str]]],
) -> dict[str, object]:
    metadata = extract_metadata(tables)
    # Fallback: extract metadata from paragraphs if table-based extraction is thin
    if len(metadata) < 3:
        para_meta = extract_metadata_from_paragraphs(paragraphs)
        for key, value in para_meta.items():
            metadata.setdefault(key, value)

    assignment_text = collect_assignment_text(tables)
    # Prefer paragraph text when it's significantly richer — common for PDFs
    # where assignment instructions appear as plain text rather than in tables
    full_para_text = " ".join(paragraphs)
    if len(full_para_text) > len(assignment_text) * 2:
        assignment_text = full_para_text
    elif len(assignment_text) < 200 and full_para_text:
        assignment_text = full_para_text

    task_sections = extract_task_sections(assignment_text)
    rubric_criteria = extract_weighted_criteria(tables)
    # Fallback: extract from grading criteria table
    if not rubric_criteria:
        rubric_criteria = extract_grading_criteria_from_table(tables)

    learning_outcomes = extract_learning_outcomes(tables)
    # Fallback: extract LOs from paragraphs
    if not learning_outcomes:
        learning_outcomes = extract_los_from_paragraphs(paragraphs)

    title = extract_assignment_title(paragraphs, assignment_text)
    module = metadata.get("Module", "")
    due_date = (
        metadata.get("To be submitted on", "")
        or metadata.get("To be submitted on:", "")
        or metadata.get("Submission Date", "")
    )
    submission_method = metadata.get("Submission Method", "")
    word_limit = metadata.get("Length", "") or metadata.get("Word Limit", "")

    summary = {
        "source_file": source_file.name,
        "title": title,
        "module": module,
        "assessment_type": infer_assignment_type(metadata, assignment_text),
        "due_date": due_date,
        "submission_method": submission_method,
        "word_limit": word_limit,
        "task_sections": task_sections,
        "rubric_criteria": rubric_criteria,
        "learning_outcomes": learning_outcomes,
    }
    return summary


def build_manifest(output_dir: Path) -> dict[str, object]:
    return {
        "output_dir": str(output_dir),
        "files": [
            {"name": "01_brief_extract.md", "purpose": "Human-readable extract of the input Word brief."},
            {"name": "02_brief_summary.json", "purpose": "Structured metadata, task sections, and rubric signals."},
            {"name": "03_agent_workflow.md", "purpose": "Agent roles, handoffs, and done criteria for the assignment run."},
            {"name": "04_research_plan.md", "purpose": "Task-to-evidence matrix for source collection and analysis."},
            {"name": "05_assignment_scaffold.md", "purpose": "Draft-ready section outline aligned to the brief."},
            {"name": "06_execution_checklist.md", "purpose": "Pre-submission checklist for drafting and QA."},
            {"name": "07_draft_starter.md", "purpose": "Editable draft starter with placeholders for claims, evidence, and citations."},
            {"name": "08_figure_table_plan.md", "purpose": "Suggested visuals and tables based on the brief structure."},
            {"name": "09_harvard_reference_guide.md", "purpose": "Harvard-style citation and reference templates."},
            {"name": "10_artifact_manifest.json", "purpose": "Machine-readable list of generated workflow artifacts."},
        ],
    }


def _add_md_section_to_docx(doc: Document, md_text: str) -> None:
    """Parse simple markdown text and add it to a python-docx Document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for line in md_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        # Headings
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        # Numbered items
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s*", "", stripped)
            p = doc.add_paragraph(text, style="List Number")
        # Table rows (markdown tables)
        elif stripped.startswith("|") and stripped.endswith("|"):
            # Skip separator rows like |---|---|
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                continue
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            # Create a table for the first row, or add rows to existing
            if not hasattr(doc, "_last_md_table") or doc._last_md_table is None:
                tbl = doc.add_table(rows=1, cols=len(cells))
                tbl.style = "Light Grid Accent 1"
                for i, cell in enumerate(cells):
                    tbl.rows[0].cells[i].text = cell
                doc._last_md_table = tbl  # type: ignore[attr-defined]
            else:
                tbl = doc._last_md_table
                row = tbl.add_row()
                for i, cell in enumerate(cells):
                    if i < len(row.cells):
                        row.cells[i].text = cell
        # Backtick code blocks — just add as plain text
        elif stripped.startswith("`") and stripped.endswith("`"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("`"))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        # Regular paragraph
        else:
            doc.add_paragraph(stripped)

        # Reset table tracker when we hit a non-table line
        if not (stripped.startswith("|") and stripped.endswith("|")):
            doc._last_md_table = None  # type: ignore[attr-defined]


def _render_figure(fig_data: dict[str, Any], figures_dir: Path, index: int) -> Path | None:
    """Render a matplotlib figure from structured data."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    try:
        fig_type = fig_data.get("type", "bar")
        labels = fig_data.get("labels", [])
        values = fig_data.get("values", [])
        if not labels or not values:
            return None

        output = figures_dir / f"figure_{index}.png"

        if fig_type == "horizontal_bar":
            fig, ax = plt.subplots(figsize=(8.5, 4.8))
            colors = ["#16324F", "#235789", "#C1292E", "#2E9A4A", "#D4A32A", "#7B4EA4"]
            bars = ax.barh(labels, values, color=colors[:len(labels)], height=0.6)
            ax.set_xlabel(fig_data.get("xlabel", ""), fontsize=10)
            ax.set_title(fig_data.get("title", ""), fontsize=13, weight="bold")
            ax.grid(axis="x", linestyle="--", alpha=0.35)
            ax.spines["top"].set_visible(False)
            ax.spines["right"].set_visible(False)
            ax.invert_yaxis()
            for bar, value in zip(bars, values):
                ax.text(value + 0.15, bar.get_y() + bar.get_height() / 2, str(value), va="center", fontsize=10)
        else:
            fig, ax = plt.subplots(figsize=(8.5, 4.8))
            colors = ["#16324F", "#235789", "#C1292E", "#2E9A4A", "#D4A32A", "#7B4EA4"]
            bars = ax.bar(labels, values, color=colors[:len(labels)], width=0.6)
            ax.set_ylabel(fig_data.get("ylabel", ""), fontsize=10)
            ax.set_xlabel(fig_data.get("xlabel", ""), fontsize=10)
            ax.set_title(fig_data.get("title", ""), fontsize=13, weight="bold")
            ax.grid(axis="y", linestyle="--", alpha=0.35)
            ax.spines["top"].set_visible(False)
            ax.spines["right"].set_visible(False)
            for bar, value in zip(bars, values):
                ax.text(bar.get_x() + bar.get_width() / 2, value + 0.3, str(value), ha="center", fontsize=10)

        fig.tight_layout()
        fig.savefig(output, dpi=220, bbox_inches="tight")
        plt.close(fig)
        return output
    except Exception:
        logger.warning("Failed to render figure", exc_info=True)
        return None


def add_caption_fn(doc: Document, text: str) -> None:
    """Add a centred italic caption paragraph to a docx."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def _add_table_to_docx(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Add a formatted table to a docx document."""
    from docx.shared import Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if not headers or not rows:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                run.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            if i < len(cells):
                cells[i].text = val
                for paragraph in cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
                cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_combined_docx(summary: dict[str, object], source_file: Path, content: dict[str, Any] | None = None) -> Document:
    """Build a single Word document containing the actual assignment draft content."""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Title page ──
    # Add some blank space
    for _ in range(4):
        doc.add_paragraph("")

    title_para = doc.add_heading(str(summary.get("title", "Assignment")), level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if summary.get("module"):
        run = subtitle.add_run(f"Module: {summary['module']}\n")
        run.font.size = Pt(14)
    if summary.get("assessment_type"):
        run = subtitle.add_run(f"Assessment Type: {summary['assessment_type']}\n")
        run.font.size = Pt(12)
    if summary.get("word_limit"):
        run = subtitle.add_run(f"Word Limit: {summary['word_limit']}\n")
        run.font.size = Pt(12)
    if summary.get("due_date"):
        run = subtitle.add_run(f"Submission Date: {summary['due_date']}\n")
        run.font.size = Pt(12)

    student_info = doc.add_paragraph()
    student_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = student_info.add_run("\nStudent Name: [Your Name]\nStudent ID: [Your ID]")
    run.font.size = Pt(12)

    doc.add_page_break()

    # ── Word count allocation (if available) ──
    word_limit = str(summary.get("word_limit", ""))
    task_sections = summary.get("task_sections", [])  # type: ignore[assignment]
    rubric_criteria = summary.get("rubric_criteria", [])  # type: ignore[assignment]
    learning_outcomes = summary.get("learning_outcomes", [])  # type: ignore[assignment]

    budget = word_count_budget(word_limit, rubric_criteria, len(task_sections))  # type: ignore[arg-type]

    if budget:
        doc.add_heading("Indicative Word-Count Allocation", level=2)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Light Grid Accent 1"
        tbl.rows[0].cells[0].text = "Section"
        tbl.rows[0].cells[1].text = "Approx. Words"
        for section_name, word_count in budget:
            row = tbl.add_row()
            row.cells[0].text = section_name
            row.cells[1].text = f"~{word_count}"
        doc.add_paragraph("")

    # ── Learning outcomes ──
    if learning_outcomes:
        doc.add_heading("Learning Outcomes Addressed", level=2)
        for lo in learning_outcomes:  # type: ignore[union-attr]
            doc.add_paragraph(str(lo), style="List Bullet")
        doc.add_paragraph("")

    doc.add_page_break()

    # ── Introduction ──
    doc.add_heading("Introduction", level=1)
    if content and content.get("introduction"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.add_run(content["introduction"])
    else:
        intro_guide = doc.add_paragraph()
        run = intro_guide.add_run(
            "[Frame the assignment topic, define the scope, state the research problem/questions, "
            "and outline the structure of the paper. Include at least one verified citation to establish context.]"
        )
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.font.size = Pt(11)
    doc.add_paragraph("")

    # ── Each detected section as a real draft section ──
    llm_sections = {}
    llm_tables = {}
    llm_figures = {}
    if content:
        for sec in content.get("sections", []):
            llm_sections[sec["title"]] = sec["body"]
        for tbl in content.get("tables", []):
            llm_tables.setdefault(tbl["section"], []).append(tbl)
        for fig in content.get("figures", []):
            llm_figures.setdefault(fig["section"], []).append(fig)

    figures_dir = source_file.parent / "figures" if content else None
    if figures_dir:
        figures_dir.mkdir(parents=True, exist_ok=True)

    for section in task_sections:  # type: ignore[union-attr]
        section_title = str(section.get("title", "Untitled"))  # type: ignore[union-attr]
        section_desc = str(section.get("description", ""))  # type: ignore[union-attr]

        # Skip if the section is basically the intro or conclusion (we add those separately)
        title_lower = section_title.lower()
        if title_lower in ("introduction", "conclusion"):
            continue

        doc.add_heading(section_title, level=1)

        # ── LLM-generated body ──
        llm_body = llm_sections.get(section_title)
        if llm_body:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
            p.add_run(llm_body)
            doc.add_paragraph("")

            # Insert LLM-generated tables for this section
            for tbl_data in llm_tables.get(section_title, []):
                add_caption_fn(doc, tbl_data.get("caption", ""))
                _add_table_to_docx(doc, tbl_data.get("headers", []), tbl_data.get("rows", []))
                doc.add_paragraph("")

            # Insert LLM-generated figures
            fig_idx = 0
            for fig_data in llm_figures.get(section_title, []):
                fig_idx += 1
                fig_path = _render_figure(fig_data, figures_dir, fig_idx)  # type: ignore[arg-type]
                if fig_path and fig_path.exists():
                    doc.add_picture(str(fig_path), width=Inches(5.5))
                add_caption_fn(doc, fig_data.get("caption", ""))
                doc.add_paragraph("")
            continue

        # ── Fallback: placeholder guidance (no LLM) ──
        # Brief requirement as guidance note
        if section_desc:
            guide_para = doc.add_paragraph()
            run = guide_para.add_run(f"Brief requirement: {section_desc}")
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            run.font.size = Pt(10)
            run.italic = True
            doc.add_paragraph("")

        # Sub-task headings (when detected)
        sub_tasks_raw = section.get("sub_tasks", [])  # type: ignore[union-attr]
        if sub_tasks_raw:
            for idx, task in enumerate(sub_tasks_raw, 1):  # type: ignore[union-attr]
                heading = task.strip()
                if heading and heading[0].islower():
                    heading = heading[0].upper() + heading[1:]
                doc.add_heading(f"{idx}. {heading}", level=2)
                note = doc.add_paragraph()
                run = note.add_run("[Answer this sub-task directly. Apply theory, cite evidence, and state your judgement.]")
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                run.font.size = Pt(10)
                doc.add_paragraph("")

        # Theory / framework placeholder
        theory_head = doc.add_paragraph()
        run = theory_head.add_run("Theory or framework to apply:")
        run.bold = True
        run.font.size = Pt(11)
        theory_placeholder = doc.add_paragraph()
        run = theory_placeholder.add_run("[Add the named model, theory, or standard you will use in this section.]")
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        doc.add_paragraph("")

        # Evidence slots
        evidence_head = doc.add_paragraph()
        run = evidence_head.add_run("Evidence to insert:")
        run.bold = True
        run.font.size = Pt(11)
        for i in range(1, 4):
            ev = doc.add_paragraph(style="List Bullet")
            run = ev.add_run(f"[Source {i}: what it proves — (Author, Year)]")
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        doc.add_paragraph("")

        # Draft paragraph plan
        plan_head = doc.add_paragraph()
        run = plan_head.add_run("Draft paragraph plan:")
        run.bold = True
        run.font.size = Pt(11)

        para_plans = [
            "Set context and scope.",
            "Apply theory or framework.",
            "Critically evaluate the evidence.",
            "State implications, comparison, or recommendation.",
        ]
        for j, plan in enumerate(para_plans, 1):
            pp = doc.add_paragraph(style="List Number")
            run = pp.add_run(f"[{plan}]")
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # Source recommendation
        sources = recommend_sources(section_title)
        if sources:
            src_head = doc.add_paragraph()
            run = src_head.add_run("Recommended sources: ")
            run.bold = True
            run.font.size = Pt(10)
            run = src_head.add_run(sources)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        doc.add_paragraph("")

    # ── Conclusion ──
    doc.add_heading("Conclusion", level=1)
    if content and content.get("conclusion"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.add_run(content["conclusion"])
    else:
        concl_guide = doc.add_paragraph()
        run = concl_guide.add_run(
            "[Summarise the key findings, highlight the contribution of the study, "
            "and suggest future research directions. Do not introduce new material here.]"
        )
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.font.size = Pt(11)
    doc.add_paragraph("")

    doc.add_page_break()

    # ── Reference List ──
    doc.add_heading("Reference List", level=1)

    llm_refs = content.get("references", []) if content else []
    if llm_refs:
        for ref in llm_refs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(ref.get("citation", ""))
            run.font.size = Pt(11)
            note = ref.get("relevance", "")
            if note:
                run2 = p.add_run(f"  [{note}]")
                run2.font.size = Pt(9)
                run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                run2.italic = True
    else:
        ref_guide = doc.add_paragraph()
        run = ref_guide.add_run("Harvard Referencing Format — add only sources you actually used and verified.")
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.italic = True
        doc.add_paragraph("")

        # Add reference format examples
        ref_templates = [
            ("Journal Article", "Author, A.A. and Author, B.B. (Year) 'Article title', Journal Title, volume(issue), pp. xx-xx."),
            ("Book", "Author, A.A. (Year) Book Title. Place of publication: Publisher."),
            ("Book Chapter", "Author, A.A. (Year) 'Chapter title', in Editor, E.E. (ed.) Book Title. Publisher, pp. xx-xx."),
            ("Web / Report", "Organisation (Year) Title. Available at: URL (Accessed: Date)."),
        ]
        for ref_type, template in ref_templates:
            ref_p = doc.add_paragraph()
            run = ref_p.add_run(f"[{ref_type}]: ")
            run.bold = True
            run.font.size = Pt(10)
            run = ref_p.add_run(template)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph("")
    doc.add_page_break()

    # ── Rubric Self-Audit ──
    doc.add_heading("Rubric Self-Audit", level=1)
    audit_intro = doc.add_paragraph()
    run = audit_intro.add_run("Check each criterion before submission:")
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph("")

    if rubric_criteria:
        for crit in rubric_criteria:  # type: ignore[union-attr]
            p = doc.add_paragraph(style="List Bullet")
            crit_name = str(crit.get("criterion", ""))  # type: ignore[union-attr]
            crit_weight = str(crit.get("weight", ""))  # type: ignore[union-attr]
            run = p.add_run(f"☐ {crit_name} ({crit_weight}): ")
            run.bold = True
            run = p.add_run("[Note where and how you address this in your paper.]")
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    else:
        doc.add_paragraph("☐ [Check rubric criteria and confirm each is addressed.]", style="List Bullet")

    doc.add_paragraph("")

    # ── Final submission checks ──
    doc.add_heading("Final Checks Before Submission", level=2)
    checks = [
        "Every citation matches a real, verified source.",
        "Every claim is supported by evidence.",
        "All placeholder text has been replaced with your own writing.",
        "Word count is within the permitted range.",
        "Submission format matches the brief requirements.",
        "Turnitin similarity is acceptable.",
        "The assignment reflects your own judgement and critical analysis.",
    ]
    for check in checks:
        doc.add_paragraph(f"☐ {check}", style="List Bullet")

    return doc


def build_output_summary_markdown(summary: dict[str, object], source_file: Path, docx_name: str) -> str:
    """Build a single markdown summary of what was generated."""
    lines: list[str] = [
        f"# Assignment Co-Pilot Output Summary",
        "",
        f"**Source brief:** `{source_file.name}`",
        f"**Generated file:** `{docx_name}`",
        "",
        "---",
        "",
        "## Brief Metadata",
        "",
        f"| Field | Value |",
        f"| --- | --- |",
        f"| Title | {summary.get('title', 'N/A')} |",
        f"| Module | {summary.get('module', 'N/A')} |",
        f"| Assessment Type | {summary.get('assessment_type', 'N/A')} |",
        f"| Word Limit | {summary.get('word_limit', 'N/A')} |",
        f"| Due Date | {summary.get('due_date', 'N/A')} |",
        "",
        "## Sections Detected",
        "",
    ]

    task_sections = summary.get("task_sections", [])
    for i, sec in enumerate(task_sections, 1):  # type: ignore[arg-type]
        desc = sec.get("description", "")  # type: ignore[union-attr]
        title = sec.get("title", "Untitled")  # type: ignore[union-attr]
        lines.append(f"{i}. **{title}**" + (f" — {desc[:80]}..." if len(desc) > 80 else (f" — {desc}" if desc else "")))
    lines.append("")

    learning_outcomes = summary.get("learning_outcomes", [])
    if learning_outcomes:
        lines.append("## Learning Outcomes")
        lines.append("")
        for lo in learning_outcomes:  # type: ignore[union-attr]
            lines.append(f"- {lo}")
        lines.append("")

    rubric_criteria = summary.get("rubric_criteria", [])
    if rubric_criteria:
        lines.append("## Rubric Criteria")
        lines.append("")
        for crit in rubric_criteria:  # type: ignore[union-attr]
            lines.append(f"- **{crit['criterion']}** ({crit['weight']})")  # type: ignore[index]
        lines.append("")

    lines.extend([
        "## What's Inside the `.docx`",
        "",
        "| Section | Description |",
        "| --- | --- |",
        "| Assignment Scaffold | Section-by-section outline with word-count allocation and LO mapping |",
        "| Execution Checklist | 13-point MSc-quality pre-submission checklist |",
        "| Research Plan | Task-to-evidence matrix for source collection |",
        "| Draft Starter | Editable draft with placeholders and rubric self-audit |",
        "| Figure & Table Plan | Suggested visuals and data tables |",
        "| Harvard Reference Guide | Citation templates for 8 source types |",
        "| Agent Workflow | Roles, handoffs, and done criteria |",
        "",
    ])

    return "\n".join(lines)


def generate_for_brief(docx_path: Path, base_output_dir: Path) -> Path:
    if not docx_path.exists():
        raise FileNotFoundError(f"Brief file not found: {docx_path}")

    suffix = docx_path.suffix.lower()
    if suffix not in (".docx", ".pdf"):
        raise ValueError(f"Expected a .docx or .pdf file, got: {docx_path.suffix}")

    if suffix == ".pdf":
        paragraphs, tables = extract_pdf_document(docx_path)
    else:
        paragraphs, tables = extract_document(docx_path)

    if not tables:
        raise ValueError(f"No tables found in {docx_path.name}; cannot extract assignment metadata.")

    summary = build_summary(docx_path, paragraphs, tables)
    logger.info("Built summary for %s: %d task sections, %d rubric criteria.",
                docx_path.name, len(summary.get("task_sections", [])), len(summary.get("rubric_criteria", [])))

    target_dir = base_output_dir / docx_path.stem
    target_dir.mkdir(parents=True, exist_ok=True)

    # Clean up old artifacts (both old-style numbered and new-style)
    for pattern in ("*.md", "*.json", "*.docx"):
        for file_path in target_dir.glob(pattern):
            file_path.unlink()

    # Generate template-based content (analysis, citations, tables, figures)
    content = build_template_content(summary)

    # Build and save the single .docx
    docx_name = f"{docx_path.stem}_helper.docx"
    doc = build_combined_docx(summary, docx_path, content=content)
    doc.save(str(target_dir / docx_name))

    # Build and save the single .md summary
    md_name = f"{docx_path.stem}_summary.md"
    md_content = build_output_summary_markdown(summary, docx_path, docx_name)
    (target_dir / md_name).write_text(md_content, encoding="utf-8")

    logger.info("Generated %s and %s in %s", docx_name, md_name, target_dir)
    return target_dir


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Read assignment brief .docx files and generate agentic workflow artifacts."
    )
    parser.add_argument("briefs", nargs="+", help="Path(s) to assignment brief .docx files.")
    parser.add_argument(
        "--output-dir",
        default="workflow_runs",
        help="Directory where generated artifact folders should be written.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output_dir = Path(args.output_dir)
    created_dirs: list[Path] = []

    for brief in args.briefs:
        docx_path = Path(brief)
        created_dirs.append(generate_for_brief(docx_path, output_dir))

    for directory in created_dirs:
        print(f"Generated assignment helper package in: {directory}")


if __name__ == "__main__":
    main()
