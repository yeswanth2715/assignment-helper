from __future__ import annotations

import argparse
import re
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ARTICLE_MATRIX = [
    (
        "Carvalho et al. (2019)",
        "Systematic review",
        "General manufacturing",
        "Model choice must fit data structure and failure context.",
    ),
    (
        "Zonta et al. (2020)",
        "Systematic review",
        "Industry 4.0 manufacturing",
        "Predictive maintenance is multidisciplinary and value depends on actionability.",
    ),
    (
        "Silvestri et al. (2020)",
        "Systematic review",
        "Smart factory maintenance",
        "Technology adoption also changes roles, management routines, and organization.",
    ),
    (
        "Ruiz-Sarmiento et al. (2020)",
        "Industrial case",
        "Steel hot rolling",
        "Large-scale real factory data can support predictive maintenance modelling.",
    ),
    (
        "Cinar et al. (2020)",
        "Review",
        "Sustainable smart manufacturing",
        "Algorithm performance depends on data quality, size, and variable relevance.",
    ),
    (
        "Cardoso and Ferreira (2021)",
        "Applied study",
        "Industrial maintenance",
        "AI becomes valuable when data are translated into maintenance knowledge.",
    ),
    (
        "Abidi et al. (2022)",
        "Planning framework",
        "Sustainable manufacturing",
        "Predictive maintenance is also a planning and decision structuring problem.",
    ),
    (
        "Achouch et al. (2022)",
        "Review and workflow model",
        "Industry 4.0 maintenance",
        "Workflow maturity is constrained by financial, data, and deployment barriers.",
    ),
    (
        "Farahani et al. (2022)",
        "Industrial case",
        "Injection molding",
        "Cloud-edge integration improves anomaly detection and process visibility.",
    ),
    (
        "Lamban et al. (2022)",
        "Cyber-physical system case",
        "Machine-tool environment",
        "KPI dashboards convert prediction into productivity and cost benefits.",
    ),
    (
        "Meddaoui et al. (2023)",
        "Manufacturing case study",
        "Production process",
        "Timely intervention improves quality and failure prediction reliability.",
    ),
    (
        "Benhanifia et al. (2025)",
        "Systematic review",
        "Manufacturing sector",
        "Operational value is clear, but ROI measurement and adoption remain uneven.",
    ),
]


REVIEW_PROTOCOL = [
    ("Time window", "2019-2025"),
    ("Source types", "Peer-reviewed journal articles"),
    ("Search focus", "Predictive maintenance, AI/ML, Industry 4.0, manufacturing, industrial prognosis"),
    ("Databases / outlets", "ScienceDirect, MDPI, SpringerLink, related academic indexes"),
    ("Inclusion logic", "Manufacturing relevance, predictive-maintenance depth, AI/digital capability, engineering-management usefulness"),
    ("Exclusion logic", "Non-journal pieces, non-manufacturing focus, purely theoretical maintenance discussion"),
    ("Analytical sample", "12 core journal articles plus 2 contextual barrier/trend studies"),
    ("Coding dimensions", "Research type, industrial context, enabling capability, value signal, implementation barrier"),
]


HYPOTHESIS_TABLE = [
    ("H1", "AI-enabled predictive maintenance improves uptime, cost efficiency, and production quality.", "10 of 12 studies", "Strong support"),
    ("H2", "Data integration, sensing capability, and decision visibility strengthen realized gains.", "9 of 12 studies", "Supported"),
    ("H3", "Strategy, ROI, cybersecurity, and legacy-system barriers constrain realized value.", "8 of 12 studies", "Strong support"),
]


BARRIER_RESPONSE = [
    ("Weak digital strategy", "Fragmented pilots and unclear scaling path", "Set a plant-level maintenance digital roadmap tied to business KPIs"),
    ("ROI uncertainty", "Projects stall after proof-of-concept", "Use phased benefits tracking: downtime, scrap, maintenance cost, schedule adherence"),
    ("Legacy-system integration", "Data silos and low data reliability", "Adopt retrofit sensing, middleware, and staged interoperability architecture"),
    ("Cybersecurity exposure", "Low trust in connected maintenance systems", "Embed security governance, access controls, and incident ownership early"),
    ("Data scarcity / poor data quality", "Models underperform or drift", "Improve data governance, feature discipline, and failure labelling routines"),
    ("Workforce readiness", "Predictions do not change real interventions", "Train engineers and planners to act on predictive signals and dashboards"),
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build a polished final assignment .docx with figures and tables from a markdown draft."
    )
    parser.add_argument("markdown_file", help="Path to the final markdown draft.")
    parser.add_argument(
        "--output",
        help="Optional .docx output path. Defaults beside the markdown file.",
    )
    return parser.parse_args()


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name, size in [("Title", 18), ("Heading 1", 14), ("Heading 2", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def parse_markdown(markdown_path: Path) -> tuple[str, dict[str, str], list[tuple[str, list[str]]]]:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    title = ""
    metadata: dict[str, str] = {}
    sections: list[tuple[str, list[str]]] = []

    current_heading: str | None = None
    current_lines: list[str] = []

    for line in lines:
        if line.startswith("# ") and not title:
            title = line[2:].strip()
            continue

        meta_match = re.match(r"^\*\*(.+?):\*\*\s*(.+?)\s*$", line)
        if meta_match and current_heading is None:
            metadata[meta_match.group(1).strip()] = meta_match.group(2).strip()
            continue

        if line.startswith("## "):
            if current_heading is not None:
                sections.append((current_heading, current_lines))
            current_heading = line[3:].strip()
            current_lines = []
            continue

        if current_heading is not None:
            current_lines.append(line)

    if current_heading is not None:
        sections.append((current_heading, current_lines))

    return title, metadata, sections


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def add_source_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def style_cell_text(cell, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.bold = bold or run.bold
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        style_cell_text(hdr[idx], bold=True)
        if widths:
            hdr[idx].width = Inches(widths[idx])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            style_cell_text(cells[idx])
            if widths:
                cells[idx].width = Inches(widths[idx])

    doc.add_paragraph("")


def add_markdown_content(doc: Document, lines: list[str]) -> None:
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines):
                candidate = lines[i].strip()
                if candidate.startswith("|") and candidate.endswith("|"):
                    table_lines.append(candidate)
                    i += 1
                else:
                    break

            parsed_rows = []
            for tbl_line in table_lines:
                if re.match(r"^\|[\s\-:|]+\|$", tbl_line):
                    continue
                parsed_rows.append([cell.strip() for cell in tbl_line.strip("|").split("|")])

            if parsed_rows:
                headers = parsed_rows[0]
                rows = [tuple(row) for row in parsed_rows[1:]]
                add_table(doc, headers, rows)
            continue

        if stripped.startswith("![") and "](" in stripped:
            match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
            if match:
                image_path = Path(match.group(2))
                if image_path.exists():
                    doc.add_picture(str(image_path), width=Inches(6.0))
                    add_caption(doc, match.group(1))
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(stripped[2:])
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line:
                break
            if next_line.startswith(("### ", "|", "- ", "![", "## ")):
                break
            paragraph_lines.append(next_line)
            i += 1

        paragraph_text = " ".join(paragraph_lines)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.add_run(paragraph_text)


def create_hypothesis_figure(figures_dir: Path) -> Path:
    output = figures_dir / "figure_1_hypothesis_support.png"
    labels = ["H1: Performance", "H2: Integration", "H3: Barriers"]
    values = [10, 9, 8]
    colors = ["#16324F", "#235789", "#C1292E"]

    fig, ax = plt.subplots(figsize=(8.5, 4.8))
    bars = ax.barh(labels, values, color=colors, height=0.6)
    ax.set_xlim(0, 12)
    ax.set_xlabel("Number of reviewed studies providing aligned evidence", fontsize=10)
    ax.set_title("Evidential Support Across the Three Hypotheses", fontsize=13, weight="bold")
    ax.grid(axis="x", linestyle="--", alpha=0.35)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.invert_yaxis()

    for bar, value in zip(bars, values):
        ax.text(value + 0.15, bar.get_y() + bar.get_height() / 2, str(value), va="center", fontsize=10)

    fig.tight_layout()
    fig.savefig(output, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return output


def create_capability_figure(figures_dir: Path) -> Path:
    output = figures_dir / "figure_2_capability_model.png"
    fig, ax = plt.subplots(figsize=(10.5, 4.8))
    ax.axis("off")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4)

    boxes = [
        ((0.45, 1.3), 2.2, 1.2, "#DCEAF7", "Enablers\nSensors\nData fusion\nCPS / cloud-edge"),
        ((3.2, 1.3), 2.2, 1.2, "#E8F3E6", "Predictive Maintenance\nDetection\nForecasting\nPlanning"),
        ((5.95, 1.3), 2.2, 1.2, "#FBE6D4", "Decision Layer\nDashboards\nScheduling\nResource timing"),
        ((8.7, 1.3), 1.0, 1.2, "#F8D7DA", "Outcomes\nUptime\nQuality\nCost"),
    ]

    for (x, y), w, h, color, text in boxes:
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.03,rounding_size=0.08",
            linewidth=1.2,
            facecolor=color,
            edgecolor="#30475E",
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=10, weight="bold")

    arrows = [
        ((2.65, 1.9), (3.2, 1.9)),
        ((5.4, 1.9), (5.95, 1.9)),
        ((8.15, 1.9), (8.7, 1.9)),
    ]
    for start, end in arrows:
        ax.add_patch(FancyArrowPatch(start, end, arrowstyle="->", mutation_scale=16, linewidth=1.6, color="#30475E"))

    barrier = FancyBboxPatch(
        (2.2, 3.0),
        5.7,
        0.65,
        boxstyle="round,pad=0.03,rounding_size=0.08",
        linewidth=1.1,
        facecolor="#FFF3CD",
        edgecolor="#856404",
    )
    ax.add_patch(barrier)
    ax.text(
        5.05,
        3.32,
        "Implementation constraints: weak strategy, ROI ambiguity, legacy systems, cybersecurity, workforce readiness",
        ha="center",
        va="center",
        fontsize=9,
        weight="bold",
    )

    ax.text(5.0, 0.55, "Author synthesis based on the reviewed literature", ha="center", fontsize=9, style="italic")
    fig.tight_layout()
    fig.savefig(output, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return output


def insert_custom_elements(doc: Document, heading: str, figures_dir: Path) -> None:
    if "Literature Review" in heading:
        add_caption(doc, "Table 1. Evidence synthesis of the core reviewed studies")
        add_table(
            doc,
            ["Study", "Design", "Context", "Core contribution"],
            ARTICLE_MATRIX,
            widths=[1.7, 1.35, 1.55, 2.65],
        )
        add_source_note(doc, "Source: author synthesis from the reviewed literature.")

    elif "Methodology" in heading:
        add_caption(doc, "Table 2. Structured review protocol and coding frame")
        add_table(
            doc,
            ["Element", "Specification"],
            REVIEW_PROTOCOL,
            widths=[1.9, 4.8],
        )
        add_source_note(doc, "Source: author-designed review protocol.")

    elif "Results and Data Analysis" in heading:
        fig1 = create_hypothesis_figure(figures_dir)
        doc.add_picture(str(fig1), width=Inches(6.4))
        add_caption(
            doc,
            "Figure 1. Relative evidential support for the three hypotheses across the reviewed studies.",
        )
        add_source_note(doc, "Source: author synthesis based on coded review findings.")

        fig2 = create_capability_figure(figures_dir)
        doc.add_picture(str(fig2), width=Inches(6.6))
        add_caption(
            doc,
            "Figure 2. Capability logic linking digital enablers, predictive maintenance routines, decision support, and operational outcomes.",
        )
        add_source_note(doc, "Source: author synthesis based on the reviewed literature.")

        add_caption(doc, "Table 3. Hypothesis assessment summary")
        add_table(
            doc,
            ["Hypothesis", "Statement", "Aligned evidence", "Assessment"],
            HYPOTHESIS_TABLE,
            widths=[0.9, 3.55, 1.4, 1.15],
        )
        add_source_note(doc, "Source: author synthesis from the analytical sample.")

        add_caption(doc, "Table 4. Barrier-to-response matrix for engineering managers")
        add_table(
            doc,
            ["Barrier", "Operational consequence", "Managerial response"],
            BARRIER_RESPONSE,
            widths=[1.5, 2.65, 2.4],
        )
        add_source_note(doc, "Source: author synthesis based on Achouch et al. (2022), Raj et al. (2020), Silvestri et al. (2020), and Benhanifia et al. (2025).")


def build_document(markdown_path: Path, output_path: Path) -> Path:
    title, metadata, sections = parse_markdown(markdown_path)
    figures_dir = output_path.parent / "figures"
    figures_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    set_document_defaults(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

    # Title page
    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.line_spacing = 1.25
    for label in ["Module", "Programme", "Student Name", "Student ID", "Submission Date"]:
        value = metadata.get(label, f"[Insert {label.lower()}]")
        line = subtitle.add_run(f"{label}: {value}\n")
        line.font.size = Pt(12)

    doc.add_page_break()

    # Body
    for heading, lines in sections:
        doc.add_heading(heading, level=1)
        insert_custom_elements(doc, heading, figures_dir)
        add_markdown_content(doc, lines)

        if heading != sections[-1][0]:
            paragraph = doc.add_paragraph("")
            paragraph.paragraph_format.space_after = Pt(0)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main() -> int:
    args = parse_args()
    markdown_path = Path(args.markdown_file)
    if not markdown_path.exists():
        raise FileNotFoundError(f"Markdown draft not found: {markdown_path}")

    output_path = Path(args.output) if args.output else markdown_path.with_suffix(".docx")
    build_document(markdown_path, output_path)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
